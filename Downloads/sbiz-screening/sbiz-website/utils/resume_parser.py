"""Extracts plain text from uploaded resumes (PDF, DOCX, TXT).

Everything happens in memory — uploaded files are never written to disk —
since resumes are personal data and this app runs on an ephemeral filesystem.
"""

import io
import re

import pdfplumber
from docx import Document

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_RESUME_CHARS = 12000  # keeps each resume to a sane size for the AI call


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _extract_pdf(file_bytes: bytes) -> str:
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()

    if text:
        return text

    # Some PDFs trip up pdfplumber but extract fine with pypdf - worth a retry
    # before giving up and telling the user it's probably a scanned image.
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))

    return "\n".join(parts).strip()


def _extract_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """Returns cleaned resume text, or raises ValueError with a user-facing message."""
    ext = filename.rsplit(".", 1)[1].lower()

    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    elif ext == "txt":
        text = _extract_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    text = re.sub(r"\n{3,}", "\n\n", text)

    if not text:
        raise ValueError(
            "Couldn't find any text in that file — it may be a scanned image "
            "without selectable text. Try a text-based PDF, DOCX, or TXT instead."
        )

    truncated = len(text) > MAX_RESUME_CHARS
    if truncated:
        text = text[:MAX_RESUME_CHARS]

    return text, truncated
